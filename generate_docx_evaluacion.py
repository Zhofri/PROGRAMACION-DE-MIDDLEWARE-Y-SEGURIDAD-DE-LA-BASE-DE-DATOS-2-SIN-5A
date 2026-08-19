"""
Generador del documento Word para la Evaluacion en Contacto con el Docente.
Utiliza python-docx para crear el archivo con formato APA 7ma edicion.
"""
import os, tempfile, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_normal_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = False
    font.italic = False
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = 2.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

def add_heading_plain(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.italic = False
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.italic = False
    return p

def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.bold = False
    run.font.italic = False
    return p

def add_fig(doc, img_path, caption, fig_num):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'Figura {fig_num}. ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    run = p.add_run(caption)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True
    return p

def add_ref(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.italic = False
    return p

def main():
    cwd = os.path.dirname(os.path.abspath(__file__))
    doc = Document()
    set_normal_style(doc)
    set_margins(doc)

    # ============= PORTADA =============
    for _ in range(4):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Universidad Internacional del Ecuador')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Facultad de Ciencias Tecnicas')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Carrera de Sistemas de la Informacion')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    for _ in range(3):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Evaluacion en Contacto con el Docente')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Portal Web Completo del Emprendimiento Materials Fadrell')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    for _ in range(3):
        doc.add_paragraph('')

    datos = [
        ('Estudiante:', 'Zhofri Joel Guaman Quichimbo'),
        ('Asignatura:', 'Programacion de Middleware y Seguridad de la Base de Datos 2-SIN-5A'),
        ('Docente:', 'Ing. Hector Guillermo Avalos Silva'),
    ]
    for label, value in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label} ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run = p.add_run(value)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Loja, Ecuador')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Agosto 2026')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # ============= 1. INTRODUCCION =============
    doc.add_page_break()
    add_heading_plain(doc, '1. Introduccion', level=1)

    add_body(doc,
        'El presente informe documenta el desarrollo completo del portal web para el '
        'emprendimiento Materials Fadrell, una empresa de venta de materiales de construccion '
        'ubicada en la ciudad de Loja. En este trabajo se integraron todas las secciones '
        'solicitadas: Home, Empresa, Productos, Oferta del Mes, Localizacion, Contacto, '
        'Noticias y Clima. Cada seccion fue alimentada con informacion coherente al '
        'emprendimiento y, en los casos de Productos, Ofertas y Noticias, los datos se '
        'cargan dinamicamente desde la base de datos a traves de endpoints REST desarrollados '
        'con Flask (Python).')

    add_body(doc,
        'La arquitectura del proyecto sigue el patron cliente-servidor. El frontend se '
        'construyo con HTML, CSS y jQuery, mientras que el backend fue implementado con '
        'el micro-framework Flask de Python, conectandose a una base de datos MariaDB '
        'administrada a traves de XAMPP. Adicionalmente, se implemento un modulo de '
        'subida FTP mediante la libreria ftplib y se integro la API de OpenWeatherMap '
        'para mostrar el clima en tiempo real de la ciudad de Loja.')

    # ============= 2. HOME Y EMPRESA =============
    doc.add_page_break()
    add_heading_plain(doc, '2. Opciones Home y Empresa', level=1)

    add_heading_plain(doc, '2.1 Seccion Home', level=2)
    add_body(doc,
        'La pagina principal del sitio web fue disenada para dar la bienvenida al visitante '
        'y presentar los valores diferenciadores de Materials Fadrell. Se incluyo un texto '
        'de bienvenida que describe la trayectoria de la empresa en Loja, asi como tres '
        'tarjetas informativas que destacan los principales beneficios: Calidad Garantizada, '
        'Entrega a Domicilio y Precios Competitivos. Cada tarjeta cuenta con un icono '
        'representativo de FontAwesome y una breve descripcion.')

    add_heading_plain(doc, '2.2 Seccion Empresa', level=2)
    add_body(doc,
        'En la seccion Empresa se incorporo toda la informacion institucional del '
        'emprendimiento. Se redactaron la Mision y la Vision de la empresa, se definieron '
        'cuatro valores corporativos (Honestidad, Responsabilidad, Calidad y Compromiso '
        'con el cliente), se diseno un organigrama empresarial que muestra la estructura '
        'jerarquica del equipo, y se incluyeron tarjetas de presentacion de los '
        'colaboradores con sus respectivos cargos y una breve descripcion de sus funciones.')

    add_body(doc,
        'El organigrama fue implementado completamente con CSS, utilizando cajas conectadas '
        'por lineas verticales y horizontales que representan la relacion jerarquica entre '
        'el Gerente General (Carlos Fadrell) y los tres jefes de area: Ventas, Logistica '
        'y Administracion. Esta solucion evita depender de imagenes estaticas y permite '
        'que la estructura se adapte al tamano de la pantalla.')

    # ============= 3. PRODUCTOS Y OFERTAS =============
    doc.add_page_break()
    add_heading_plain(doc, '3. Productos y Ofertas del Mes', level=1)

    add_heading_plain(doc, '3.1 Catalogo de Productos (carga dinamica)', level=2)
    add_body(doc,
        'Una de las mejoras mas importantes respecto a las versiones anteriores del sitio '
        'fue la conversion del catalogo de productos de estatico a dinamico. Ahora, cuando '
        'el usuario hace clic en la pestana Productos, el frontend ejecuta una peticion '
        'fetch() al endpoint /api/productos del backend Flask. Este endpoint consulta la '
        'tabla productos en la base de datos (realizando un LEFT JOIN con la tabla categorias '
        'para obtener el nombre de la categoria) y devuelve los resultados en formato JSON.')

    add_body(doc,
        'El codigo JavaScript recibe esta respuesta y construye dinamicamente las tarjetas '
        'de producto, mostrando la imagen del producto (almacenada en la carpeta assets), '
        'el nombre, la categoria, el precio y el stock disponible. Tambien se generan '
        'automaticamente los botones de filtro por categoria a partir de los datos recibidos '
        'de la API, de modo que si se agregan nuevas categorias a la base de datos, estas '
        'aparecen sin necesidad de modificar el codigo del frontend.')

    add_body(doc, 'A continuacion se muestra el codigo del endpoint en Flask:', indent=False)
    add_code(doc,
        "@app.route('/api/productos', methods=['GET'])\n"
        "def obtener_productos():\n"
        "    conexion = pymysql.connect(**DB_CONFIG)\n"
        "    with conexion.cursor() as cursor:\n"
        "        sql = 'SELECT p.id, p.nombre, p.descripcion, p.precio,\\n'\n"
        "              '       p.stock, c.nombre as categoria, p.imagen_url\\n'\n"
        "              'FROM productos p\\n'\n"
        "              'LEFT JOIN categorias c ON p.categoria_id = c.id'\n"
        "        cursor.execute(sql)\n"
        "        productos = cursor.fetchall()\n"
        "    conexion.close()\n"
        "    return jsonify(productos), 200")

    add_heading_plain(doc, '3.2 Ofertas del Mes (carga dinamica)', level=2)
    add_body(doc,
        'De manera similar, la seccion de Ofertas del Mes tambien se alimenta desde la '
        'base de datos. El endpoint /api/ofertas filtra las ofertas cuya fecha de vigencia '
        'incluye la fecha actual (usando la clausula WHERE CURRENT_DATE BETWEEN '
        'fecha_inicio AND fecha_fin). Para cada oferta activa, se calcula el precio con '
        'descuento directamente en la consulta SQL y se devuelve junto con la imagen del '
        'producto, el porcentaje de descuento y la fecha limite de la promocion.')

    add_body(doc,
        'En el frontend, cada oferta se presenta como una tarjeta con la imagen del '
        'producto, el precio original tachado, el precio con descuento destacado en color, '
        'una etiqueta con el porcentaje de descuento y la fecha de vencimiento. Si no '
        'existen ofertas activas para el mes en curso, se muestra un mensaje amigable '
        'indicando que no hay promociones disponibles por el momento.')

    # ============= 4. LOCALIZACION Y CONTACTO =============
    doc.add_page_break()
    add_heading_plain(doc, '4. Localizacion y Contacto', level=1)

    add_heading_plain(doc, '4.1 Localizacion con Google Maps', level=2)
    add_body(doc,
        'La seccion de Localizacion presenta la ubicacion fisica de Materials Fadrell '
        'mediante un mapa interactivo de Google Maps incrustado a traves de un iframe. '
        'Se incluyen los datos de contacto de la empresa: direccion completa, numero de '
        'telefono convencional, celular y correo electronico. Ademas, se proporciona un '
        'punto de referencia para facilitar la llegada al local. El diseno utiliza una '
        'cuadricula de dos columnas: la informacion de contacto a la izquierda y el mapa '
        'interactivo a la derecha.')

    add_heading_plain(doc, '4.2 Formulario de Contacto', level=2)
    add_body(doc,
        'El formulario de contacto permite a los usuarios enviar un mensaje con los '
        'campos Nombre, Email y Mensaje, junto con un boton Guardar. Al presionar el '
        'boton, JavaScript recopila los datos y los envia al endpoint /api/contacto '
        'del backend Flask mediante una peticion POST con formato JSON. El backend '
        'inserta el registro en la tabla contacto de la base de datos y, a continuacion, '
        'envia un correo electronico de acuse de recibido al remitente usando la libreria '
        'smtplib de Python con el servidor SMTP de Gmail.')

    add_body(doc,
        'Este mecanismo de acuse de recibido fue implementado con la clase EmailMessage '
        'y la conexion SMTP_SSL en el puerto 465. El correo contiene un mensaje '
        'personalizado con el nombre del usuario, confirmando que su consulta fue recibida '
        'correctamente.')

    # ============= 5. NOTICIAS Y CLIMA =============
    doc.add_page_break()
    add_heading_plain(doc, '5. Noticias y Clima', level=1)

    add_heading_plain(doc, '5.1 Seccion de Noticias (carga dinamica)', level=2)
    add_body(doc,
        'Para la seccion de Noticias se diseno una tabla en la base de datos con los '
        'campos requeridos: fecha, titulo, contenido e imagen. Al acceder a esta seccion, '
        'el frontend realiza una peticion fetch() al endpoint /api/noticias, que devuelve '
        'todas las noticias ordenadas por fecha descendente.')

    add_body(doc, 'La tabla de noticias fue creada con el siguiente script SQL:', indent=False)
    add_code(doc,
        "CREATE TABLE IF NOT EXISTS noticias (\n"
        "    id INT AUTO_INCREMENT PRIMARY KEY,\n"
        "    fecha DATE NOT NULL,\n"
        "    titulo VARCHAR(200) NOT NULL,\n"
        "    contenido TEXT NOT NULL,\n"
        "    imagen_url VARCHAR(255) DEFAULT NULL\n"
        ");")

    add_body(doc,
        'En el frontend, cada noticia se muestra con un diseno horizontal que ubica la '
        'imagen a la izquierda y los datos textuales (fecha, titulo y contenido) a la '
        'derecha, tal como se indica en el mockup proporcionado por el docente.')

    add_heading_plain(doc, '5.2 Reporteador de Clima (API)', level=2)
    add_body(doc,
        'La seccion de Clima consume la API de OpenWeatherMap para mostrar los datos '
        'meteorologicos actuales de la ciudad de Loja, Ecuador. Se utiliza el endpoint '
        'weather de la API con los parametros de consulta correspondientes (ciudad, '
        'unidades metricas e idioma espanol). La respuesta JSON incluye la temperatura '
        'actual, la humedad relativa, la descripcion del clima y la velocidad del viento.')

    add_body(doc,
        'El frontend presenta estos datos en un panel con cuatro indicadores: Temperatura '
        'Actual, Humedad, Descripcion del clima y Viento. Ademas, se muestra el icono '
        'oficial que provee la misma API para representar visualmente las condiciones '
        'climaticas actuales.')

    # ============= 6. FTP Y GITLAB =============
    doc.add_page_break()
    add_heading_plain(doc, '6. Subida FTP y Repositorio GitLab', level=1)

    add_heading_plain(doc, '6.1 Subida a FTP con ftplib', level=2)
    add_body(doc,
        'Se desarrollo un script en Python (subir_ftp.py) que utiliza la libreria ftplib '
        'para comprimir los archivos fuente del proyecto en un archivo ZIP y subirlos a '
        'un servidor FTP remoto de forma automatizada. Este proceso se realiza en tres '
        'pasos: primero se crea el archivo ZIP con los codigos fuente, luego se establece '
        'la conexion con el servidor FTP, y finalmente se sube el archivo al directorio '
        'indicado.')

    add_heading_plain(doc, '6.2 Repositorio GitLab', level=2)
    add_body(doc,
        'El codigo fuente completo del proyecto se encuentra publicado en un repositorio '
        'publico dentro de la plataforma GitLab. El repositorio contiene todos los archivos '
        'del portal web (HTML, CSS, JavaScript), el backend en Python (Flask), los scripts '
        'SQL y el archivo README con las instrucciones de despliegue.')

    add_body(doc, 'El enlace de acceso al repositorio es el siguiente:', indent=False)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run('https://gitlab.com/Zhofri/autonomo-2')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.font.bold = False

    # ============= 7. CONCLUSIONES =============
    doc.add_page_break()
    add_heading_plain(doc, '7. Conclusiones', level=1)

    add_body(doc,
        'Con el desarrollo de este proyecto se logro construir un portal web completo '
        'para el emprendimiento Materials Fadrell, integrando ocho secciones funcionales '
        'que cubren desde la presentacion institucional de la empresa hasta la consulta '
        'del clima en tiempo real. Las secciones de Productos, Ofertas del Mes y Noticias '
        'se alimentan de forma dinamica desde la base de datos, lo cual permite actualizar '
        'el contenido sin necesidad de modificar el codigo fuente del frontend.')

    add_body(doc,
        'La integracion del backend con Flask y la base de datos MariaDB permitio comprender '
        'de manera practica como funciona la comunicacion entre el cliente y el servidor '
        'mediante peticiones HTTP asincronas. El formulario de contacto, por su parte, '
        'demostro como persistir datos ingresados por el usuario y enviar notificaciones '
        'automaticas por correo electronico.')

    add_body(doc,
        'Finalmente, el uso de herramientas como ftplib para el despliegue automatizado, '
        'GitLab para el control de versiones, y la API de OpenWeatherMap para el consumo '
        'de servicios REST externos, refuerzan las competencias tecnicas adquiridas a lo '
        'largo del curso y preparan al estudiante para escenarios reales de desarrollo '
        'de software.')

    # ============= 8. REFERENCIAS =============
    add_heading_plain(doc, '8. Referencias', level=1)

    add_ref(doc,
        'Grinberg, M. (2018). Flask Web Development: Developing Web Applications with Python '
        '(2da ed.). O\'Reilly Media. https://www.oreilly.com/library/view/flask-web-development/9781491991725/')
    add_ref(doc,
        'OpenWeather Ltd. (2024). Current weather data - OpenWeatherMap API documentation. '
        'https://openweathermap.org/current')
    add_ref(doc,
        'Python Software Foundation. (2024). ftplib - FTP protocol client. Python 3.12 Documentation. '
        'https://docs.python.org/3/library/ftplib.html')
    add_ref(doc,
        'Python Software Foundation. (2024). smtplib - SMTP protocol client. Python 3.12 Documentation. '
        'https://docs.python.org/3/library/smtplib.html')
    add_ref(doc,
        'Mozilla Developer Network. (2024). Fetch API. MDN Web Docs. '
        'https://developer.mozilla.org/es/docs/Web/API/Fetch_API')
    add_ref(doc,
        'jQuery Foundation. (2024). jQuery API Documentation. https://api.jquery.com/')
    add_ref(doc,
        'Wikipedia contributors. (2024). Flask (web framework). Wikipedia, la enciclopedia libre. '
        'https://es.wikipedia.org/wiki/Flask_(framework)')
    add_ref(doc,
        'MariaDB Foundation. (2024). MariaDB Server Documentation. '
        'https://mariadb.com/kb/en/documentation/')

    # ============= GUARDAR =============
    temp_path = os.path.join(tempfile.gettempdir(), 'Guaman_Zhofri_Evaluacion_ProgramacionBD.docx')
    doc.save(temp_path)

    output_path = os.path.join(cwd, 'Guaman_Zhofri_Evaluacion_ProgramacionBD.docx')
    try:
        shutil.copy2(temp_path, output_path)
        print(f'[OK] Documento guardado en: {output_path}')
    except Exception as e:
        print(f'[AVISO] No se pudo copiar al destino ({e}). El archivo esta en: {temp_path}')

if __name__ == '__main__':
    main()
